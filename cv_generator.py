import os
import tempfile
import subprocess
import shutil
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ─── Colors ───
BLUE     = RGBColor(0x1A, 0x52, 0x76)
GRAY     = RGBColor(0x5D, 0x6D, 0x7E)
BLACK    = RGBColor(0x1C, 0x1C, 0x1C)
BLUE_HEX = "1A5276"

# ─── Section Labels ───
LABELS = {
    'en': {
        'summary':      'Professional Summary',
        'experience':   'Experience',
        'education':    'Education',
        'skills':       'Skills',
        'certificates': 'Certificates',
        'languages':    'Languages',
    },
    'ar': {
        'summary':      'الملخص المهني',
        'experience':   'الخبرات العملية',
        'education':    'التعليم',
        'skills':       'المهارات',
        'certificates': 'الشهادات والدورات',
        'languages':    'اللغات',
    }
}


# ════════════════════════════════════════
#              HELPERS
# ════════════════════════════════════════

def font(run, size, color=None, bold=False, italic=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.bold   = bold
    run.font.italic = italic


def set_rtl(para):
    """Make paragraph RTL (Arabic direction)."""
    pPr = para._p.get_or_add_pPr()
    # bidi direction
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    # right alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)


def set_rtl_run(run):
    """Mark a run as RTL."""
    rPr = run._r.get_or_add_rPr()
    rtl_el = OxmlElement('w:rtl')
    rPr.append(rtl_el)
    # complex script font
    cs = OxmlElement('w:cs')
    cs.set(qn('w:val'), 'Arial')
    rPr.append(cs)


def apply_rtl_to_para(para):
    """Apply RTL to paragraph and all its runs."""
    set_rtl(para)
    for run in para.runs:
        set_rtl_run(run)


def add_left_tab(para, position=9026):
    """Left-aligned tab (used for dates in RTL mode)."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(position))
    tabs.append(tab)
    pPr.append(tabs)


def add_right_tab(para, position=9026):
    """Right-aligned tab (used for dates in LTR mode)."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(position))
    tabs.append(tab)
    pPr.append(tabs)


def add_section_divider(doc, title, rtl=False):
    """Section title with blue bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)

    run = para.add_run(title.upper() if not rtl else title)
    font(run, 12, BLUE, bold=True)

    if rtl:
        set_rtl(para)
        set_rtl_run(run)

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), BLUE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def remove_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_job_entry(doc, title, company, date_range, bullets, rtl=False):
    """Work experience block — supports LTR and RTL."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

    if rtl:
        # In RTL: title on right, date on left
        set_rtl(p)
        r_title = p.add_run(title)
        font(r_title, 11, BLACK, bold=True)
        set_rtl_run(r_title)
        r_sep = p.add_run('    |    ')
        font(r_sep, 10, GRAY)
        r_date = p.add_run(date_range)
        font(r_date, 10, GRAY)
        set_rtl_run(r_date)
    else:
        add_right_tab(p)
        font(p.add_run(title), 11, BLACK, bold=True)
        p.add_run('\t').font.name = 'Arial'
        font(p.add_run(date_range), 10, GRAY)

    pc = doc.add_paragraph()
    pc.paragraph_format.space_before = Pt(0)
    pc.paragraph_format.space_after  = Pt(4)
    r_co = pc.add_run(company)
    font(r_co, 10, BLUE, italic=True)
    if rtl:
        set_rtl(pc)
        set_rtl_run(r_co)

    for b in bullets:
        pb = doc.add_paragraph(style='List Bullet')
        pb.paragraph_format.space_before = Pt(2)
        pb.paragraph_format.space_after  = Pt(2)
        r_b = pb.add_run(b)
        font(r_b, 10, BLACK)
        if rtl:
            set_rtl(pb)
            set_rtl_run(r_b)


def add_skills_table(doc, skills, rtl=False):
    """Two-column borderless skills table."""
    half  = (len(skills) + 1) // 2
    left  = skills[:half]
    right = skills[half:]

    table = doc.add_table(rows=half, cols=2)
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # RTL table direction
    if rtl:
        bidiVisual = OxmlElement('w:bidiVisual')
        tblPr.append(bidiVisual)

    for i in range(half):
        row = table.rows[i]

        cl = row.cells[0]
        cl.width = Twips(4513)
        remove_cell_borders(cl)
        pl = cl.paragraphs[0]
        pl.clear()
        r1 = pl.add_run("▸ ")
        font(r1, 10, BLUE)
        r2 = pl.add_run(left[i] if i < len(left) else "")
        font(r2, 10, BLACK)
        if rtl:
            set_rtl(pl)
            set_rtl_run(r1)
            set_rtl_run(r2)

        cr = row.cells[1]
        cr.width = Twips(4513)
        remove_cell_borders(cr)
        if i < len(right):
            pr = cr.paragraphs[0]
            pr.clear()
            r3 = pr.add_run("▸ ")
            font(r3, 10, BLUE)
            r4 = pr.add_run(right[i])
            font(r4, 10, BLACK)
            if rtl:
                set_rtl(pr)
                set_rtl_run(r3)
                set_rtl_run(r4)


# ════════════════════════════════════════
#           GENERATE CV (DOCX)
# ════════════════════════════════════════

def generate_cv(data: dict, lang: str = 'en') -> str:
    """
    Build a DOCX CV from user data.
    lang: 'en' (LTR) or 'ar' (RTL)
    Returns path to temp file.
    """
    rtl    = (lang == 'ar')
    labels = LABELS[lang]
    tmp_path = None

    try:
        doc = Document()

        # Page margins (A4)
        for section in doc.sections:
            section.top_margin    = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin   = Cm(1.8)
            section.right_margin  = Cm(1.8)

        # ── HEADER ──────────────────────────
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run(data['name'])
        font(r_name, 26, BLUE, bold=True)
        if rtl:
            set_rtl(p_name)
            set_rtl_run(r_name)

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(6)
        r_title = p_title.add_run(data['job_title'])
        font(r_title, 12, GRAY)
        if rtl:
            set_rtl(p_title)
            set_rtl_run(r_title)

        # Contact line
        contact_parts = [
            f"📧 {data['email']}",
            f"📞 {data['phone']}",
            f"📍 {data['location']}",
        ]
        if data.get('linkedin'):
            contact_parts.append(f"🔗 {data['linkedin']}")
        if data.get('github'):
            contact_parts.append(f"💻 {data['github']}")

        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(12)
        r_contact = p_contact.add_run("   |   ".join(contact_parts))
        font(r_contact, 9, GRAY)
        if rtl:
            set_rtl(p_contact)
            set_rtl_run(r_contact)

        # ── SUMMARY ─────────────────────────
        add_section_divider(doc, labels['summary'], rtl=rtl)
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_before = Pt(4)
        p_sum.paragraph_format.space_after  = Pt(10)
        r_sum = p_sum.add_run(data['summary'])
        font(r_sum, 10, BLACK)
        if rtl:
            set_rtl(p_sum)
            set_rtl_run(r_sum)

        # ── EXPERIENCE ──────────────────────
        add_section_divider(doc, labels['experience'], rtl=rtl)
        for exp in data['experiences']:
            add_job_entry(doc, exp['title'], exp['company'],
                          exp['date'], exp['bullets'], rtl=rtl)

        # ── EDUCATION ───────────────────────
        add_section_divider(doc, labels['education'], rtl=rtl)
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(8)
        p_edu.paragraph_format.space_after  = Pt(2)

        if rtl:
            set_rtl(p_edu)
            r_deg = p_edu.add_run(data['edu_degree'])
            font(r_deg, 11, BLACK, bold=True)
            set_rtl_run(r_deg)
            r_sep2 = p_edu.add_run('    |    ')
            font(r_sep2, 10, GRAY)
            r_date2 = p_edu.add_run(data['edu_date'])
            font(r_date2, 10, GRAY)
            set_rtl_run(r_date2)
        else:
            add_right_tab(p_edu)
            font(p_edu.add_run(data['edu_degree']), 11, BLACK, bold=True)
            p_edu.add_run('\t').font.name = 'Arial'
            font(p_edu.add_run(data['edu_date']), 10, GRAY)

        p_uni = doc.add_paragraph()
        p_uni.paragraph_format.space_before = Pt(0)
        p_uni.paragraph_format.space_after  = Pt(10)
        r_uni = p_uni.add_run(data['edu_university'])
        font(r_uni, 10, GRAY, italic=True)
        if rtl:
            set_rtl(p_uni)
            set_rtl_run(r_uni)

        # ── SKILLS ──────────────────────────
        add_section_divider(doc, labels['skills'], rtl=rtl)
        doc.add_paragraph().paragraph_format.space_before = Pt(6)
        add_skills_table(doc, data['skills'], rtl=rtl)

        # ── CERTIFICATES ────────────────────
        if data.get('certificates'):
            add_section_divider(doc, labels['certificates'], rtl=rtl)
            for cert in data['certificates']:
                p_cert = doc.add_paragraph()
                p_cert.paragraph_format.space_before = Pt(4)
                p_cert.paragraph_format.space_after  = Pt(2)
                if rtl:
                    set_rtl(p_cert)
                    r_cn = p_cert.add_run(cert['name'])
                    font(r_cn, 10, BLACK, bold=True)
                    set_rtl_run(r_cn)
                    r_sep3 = p_cert.add_run(' — ')
                    font(r_sep3, 10, GRAY)
                    r_ci = p_cert.add_run(cert['issuer'])
                    font(r_ci, 10, GRAY, italic=True)
                    set_rtl_run(r_ci)
                    r_sep4 = p_cert.add_run('    |    ')
                    font(r_sep4, 10, GRAY)
                    r_cd = p_cert.add_run(cert['date'])
                    font(r_cd, 10, GRAY)
                    set_rtl_run(r_cd)
                else:
                    add_right_tab(p_cert)
                    font(p_cert.add_run(cert['name']), 10, BLACK, bold=True)
                    font(p_cert.add_run(' — '), 10, GRAY)
                    font(p_cert.add_run(cert['issuer']), 10, GRAY, italic=True)
                    p_cert.add_run('\t').font.name = 'Arial'
                    font(p_cert.add_run(cert['date']), 10, GRAY)

        # ── LANGUAGES ───────────────────────
        add_section_divider(doc, labels['languages'], rtl=rtl)
        p_lang = doc.add_paragraph()
        p_lang.paragraph_format.space_before = Pt(6)
        r_lang = p_lang.add_run(data['languages'])
        font(r_lang, 10, BLACK)
        if rtl:
            set_rtl(p_lang)
            set_rtl_run(r_lang)

        # Save to temp file
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp_path = tmp.name
        tmp.close()
        doc.save(tmp_path)
        logger.info(f"CV generated ({lang}): {tmp_path}")
        return tmp_path

    except Exception as e:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        logger.error(f"CV generation failed: {e}")
        raise



# ════════════════════════════════════════
#        ARABIC PDF (reportlab)
# ════════════════════════════════════════

def _ar(text: str) -> str:
    """Reshape and reorder Arabic text for correct PDF display."""
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        return get_display(arabic_reshaper.reshape(str(text)))
    except Exception:
        return str(text)


def _get_arabic_font() -> tuple:
    """Find Amiri or Noto Arabic font. Returns (regular_name, bold_name)."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        ('/app/fonts/Amiri-Regular.ttf',     '/app/fonts/Amiri-Bold.ttf',     'Amiri',   'AmiriBd'),
        ('/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf',
         '/usr/share/fonts/truetype/noto/NotoNaskhArabic-Bold.ttf', 'NotoAr', 'NotoArBd'),
    ]
    for reg_path, bold_path, reg_name, bold_name in candidates:
        if os.path.exists(reg_path):
            pdfmetrics.registerFont(TTFont(reg_name, reg_path))
            if os.path.exists(bold_path):
                pdfmetrics.registerFont(TTFont(bold_name, bold_path))
            else:
                bold_name = reg_name
            return reg_name, bold_name

    return 'Helvetica', 'Helvetica-Bold'


def generate_arabic_pdf(data: dict) -> str:
    """
    Generate a proper Arabic RTL PDF using reportlab.
    Returns path to the generated PDF file.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, white
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        HRFlowable, Table, TableStyle
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_RIGHT, TA_CENTER

    C_BLUE  = HexColor('#1A5276')
    C_GRAY  = HexColor('#5D6D7E')
    C_BLACK = HexColor('#1C1C1C')

    font_reg, font_bold = _get_arabic_font()

    def S(name, size, color, align=TA_RIGHT, bold=False):
        return ParagraphStyle(
            name, fontName=font_bold if bold else font_reg,
            fontSize=size, textColor=color,
            alignment=align, leading=size * 1.5,
            wordWrap='RTL',
        )

    S_NAME    = S('name',    22, C_BLUE,  bold=True)
    S_TITLE   = S('title',   11, C_GRAY)
    S_CONTACT = S('contact',  8, C_GRAY,  align=TA_CENTER)
    S_SECTION = S('section', 11, C_BLUE,  bold=True)
    S_BODY    = S('body',    10, C_BLACK)
    S_COMPANY = S('company', 10, C_BLUE)
    S_GRAY    = S('gray',    10, C_GRAY)

    def sec(title):
        return [
            Spacer(1, 0.3 * cm),
            Paragraph(_ar(title), S_SECTION),
            HRFlowable(width='100%', thickness=1, color=C_BLUE, spaceAfter=4),
            Spacer(1, 0.1 * cm),
        ]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    tmp_path = tmp.name
    tmp.close()

    try:
        doc = SimpleDocTemplate(
            tmp_path, pagesize=A4,
            rightMargin=1.8*cm, leftMargin=1.8*cm,
            topMargin=1.5*cm,   bottomMargin=1.5*cm,
        )
        story = []

        # ── Header ──
        story.append(Paragraph(_ar(data['name']), S_NAME))
        story.append(Spacer(1, 0.1 * cm))
        story.append(Paragraph(_ar(data['job_title']), S_TITLE))
        story.append(Spacer(1, 0.2 * cm))

        contact_parts = [data['email'], data['phone'], data['location']]
        if data.get('linkedin'):
            contact_parts.append(data['linkedin'])
        if data.get('github'):
            contact_parts.append(data['github'])
        story.append(Paragraph('  |  '.join(contact_parts), S_CONTACT))
        story.append(HRFlowable(width='100%', thickness=1.5, color=C_BLUE, spaceBefore=6, spaceAfter=2))

        # ── Summary ──
        story.extend(sec(LABELS['ar']['summary']))
        story.append(Paragraph(_ar(data['summary']), S_BODY))

        # ── Experience ──
        story.extend(sec(LABELS['ar']['experience']))
        for exp in data['experiences']:
            story.append(Spacer(1, 0.15 * cm))
            line = f"{_ar(exp['date'])}    |    {_ar(exp['title'])}"
            story.append(Paragraph(line, S_BODY))
            story.append(Paragraph(_ar(exp['company']), S_COMPANY))
            for b in exp['bullets']:
                story.append(Paragraph(f"◀  {_ar(b)}", S_BODY))

        # ── Education ──
        story.extend(sec(LABELS['ar']['education']))
        story.append(Paragraph(
            f"{_ar(data['edu_date'])}    |    {_ar(data['edu_degree'])}", S_BODY
        ))
        story.append(Paragraph(_ar(data['edu_university']), S_GRAY))

        # ── Skills ──
        story.extend(sec(LABELS['ar']['skills']))
        skills = data['skills']
        half   = (len(skills) + 1) // 2
        col_w  = (A4[0] - 3.6 * cm) / 2
        rows   = []
        for i in range(half):
            l = Paragraph(f"◀  {_ar(skills[i])}", S_BODY)
            r = Paragraph(f"◀  {_ar(skills[half + i])}", S_BODY) if (half + i) < len(skills) else Paragraph('', S_BODY)
            rows.append([l, r])
        if rows:
            t = Table(rows, colWidths=[col_w, col_w])
            t.setStyle(TableStyle([
                ('ALIGN',   (0, 0), (-1, -1), 'RIGHT'),
                ('VALIGN',  (0, 0), (-1, -1), 'TOP'),
                ('GRID',    (0, 0), (-1, -1), 0, white),
                ('TOPPADDING',    (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            story.append(t)

        # ── Certificates ──
        if data.get('certificates'):
            story.extend(sec(LABELS['ar']['certificates']))
            for cert in data['certificates']:
                line = f"{_ar(cert['date'])}  |  {_ar(cert['issuer'])}  —  {_ar(cert['name'])}"
                story.append(Paragraph(line, S_BODY))

        # ── Languages ──
        story.extend(sec(LABELS['ar']['languages']))
        story.append(Paragraph(_ar(data['languages']), S_BODY))

        doc.build(story)
        logger.info(f"Arabic PDF generated: {tmp_path}")
        return tmp_path

    except Exception as e:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        logger.error(f"Arabic PDF generation failed: {e}")
        raise


# ════════════════════════════════════════
#        ENGLISH PDF (LibreOffice)
# ════════════════════════════════════════

def _find_libreoffice() -> str:
    candidates = [
        'libreoffice', 'soffice',
        '/usr/bin/libreoffice', '/usr/bin/soffice',
        '/usr/lib/libreoffice/program/soffice',
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
    ]
    for c in candidates:
        if shutil.which(c) or os.path.exists(c):
            return c
    raise FileNotFoundError("LibreOffice not found.")


def convert_to_pdf(docx_path: str) -> str:
    """Convert English DOCX to PDF via LibreOffice."""
    tmp_dir = tempfile.mkdtemp()
    try:
        lo_bin = _find_libreoffice()
        result = subprocess.run(
            [lo_bin, '--headless', '--convert-to', 'pdf',
             '--outdir', tmp_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            raise Exception(f"LibreOffice error: {result.stderr.strip()}")
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path  = os.path.join(tmp_dir, base_name + '.pdf')
        if not os.path.exists(pdf_path):
            raise Exception("PDF not found after conversion")
        return pdf_path
    except Exception:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        raise

